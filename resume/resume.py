from docx import Document

# Create a new Document
doc = Document()

# Add Name and Contact Information
doc.add_heading('Pavan Khairnar', level=1)

contact_info = (
    "Address: Row House No. 3, Behind Dhanvantri Medical College, Kamatwade, Nashik\n"
    "Email: pavansk00@gmail.com\n"
    "Phone: +91 8055719031\n"
    "LinkedIn: https://www.linkedin.com/in/pavan-khairnar-588077140/"
)
doc.add_paragraph(contact_info)

# Add Professional Summary
doc.add_heading('Professional Summary', level=2)
summary = (
    "I’m a game developer specializing in Python. I excel in creating game logic, mechanics, and APIs, working closely "
    "with game design teams to develop engaging features. My expertise includes both SQL and NoSQL databases for efficient "
    "data management, as well as building secure and scalable APIs. I also focus on real-time communication and microservices "
    "architecture, ensuring smooth and reliable operations."
)
doc.add_paragraph(summary)

# Add Skills
doc.add_heading('Skills', level=2)
skills = (
    "- Python\n"
    "- MySQL\n"
    "- Linux\n"
    "- System Design\n"
    "- AWS\n"
    "- Big Data\n"
    "- Hadoop\n"
    "- Spark\n"
    "- Power BI\n"
    "- Machine Learning"
)
doc.add_paragraph(skills)

# Add Professional Experience
doc.add_heading('Experience', level=2)

# Software Engineer Role
doc.add_heading('Software Engineer', level=3)
doc.add_paragraph('IVY Comptech Private Limited, Hyderabad | Dec 2021 - Present')
software_engineer_responsibilities = (
    "- Design, develop, and maintain game logic and mechanics using Python.\n"
    "- Collaborate with game design teams to implement game features and functionalities.\n"
    "- Develop and maintain APIs to support web and game functionalities.\n"
    "- Manage databases (SQL & NoSQL) for efficient data storage and retrieval.\n"
    "- Ensure the security, reliability, and scalability of APIs.\n"
    "- Collaborate with front-end developers for seamless web application integration.\n"
    "- Monitor and analyze logs and database data to resolve production issues.\n"
    "- Implement logging and monitoring solutions to ensure system reliability.\n"
    "- Develop and implement WebSocket-based solutions for real-time communication.\n"
    "- Design and implement microservices architecture for scalable applications."
)
doc.add_paragraph(software_engineer_responsibilities)

# POP Engineer Role
doc.add_heading('POP Engineer', level=3)
doc.add_paragraph('Tikona Infinet Private Limited, Nashik | Sept 2018 - July 2021')
pop_engineer_responsibilities = (
    "- Network monitoring, router configuration, RF device LTE link setup.\n"
    "- Network maintenance, troubleshooting, Linux machine configuration, and firewall services.\n"
    "- Provided technical support."
)
doc.add_paragraph(pop_engineer_responsibilities)

# Add Education
doc.add_heading('Education', level=2)

# Pg-Diploma
doc.add_heading('Pg-Diploma in Big Data Analytics', level=3)
doc.add_paragraph('Center for Development of Advanced Computing (CDAC), Pune | Sept 2021')

# Bachelor of Engineering
doc.add_heading('Bachelor of Engineering in Electronics & Telecommunication', level=3)
doc.add_paragraph("Pune Vidyarthi Griha's College of Engineering, Nashik | June 2017")

# Diploma in Electronics & Telecommunication
doc.add_heading('Diploma in Electronics & Telecommunication', level=3)
doc.add_paragraph("Maratha Vidya Prasarak Samaj's Rajarshi Shahu Maharaj Polytechnic, Nashik | June 2014")

# 10th School
doc.add_heading('10th School', level=3)
doc.add_paragraph('Maratha High School, Nashik | 2011')

# Add Projects
doc.add_heading('Projects', level=2)

# Socket Service
doc.add_heading('Socket Service', level=3)
socket_service_details = (
    "- Developed a scalable microservice to handle over 100,000 concurrent game players.\n"
    "- Utilized Kafka and Redis for efficient data transfer and management.\n"
    "- Designed a horizontally scalable system for adding servers to support an increasing number of players.\n"
    "- Implemented WebSocket for real-time communication."
)
doc.add_paragraph(socket_service_details)

# Kibana Integration
doc.add_heading('Kibana Integration', level=3)
kibana_integration_details = (
    "- Integrated Kibana, Logstash, and Elasticsearch for game application data tracking.\n"
    "- Created Kibana dashboards to visualize trends, identify production issues, and understand game feature functionality.\n"
    "- Analyzed player data to provide insights for business decisions and improve the gaming experience.\n"
    "- Developed comprehensive monitoring and logging solutions to enhance system reliability."
)
doc.add_paragraph(kibana_integration_details)

# Add Interests
doc.add_heading('Interests', level=2)
interests = "- Reading\n- Travel"
doc.add_paragraph(interests)

# Save the document
file_path = "Pavan_Khairnar_ATS_Resume.docx"
doc.save(file_path)

file_path
